# -*- coding: utf-8 -*-
"""
Created on Mon Mar  6 19:00:26 2023

@author: NEHA
"""


from docx import Document
from docx.shared import Inches
import pyttsx3 


def say(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()
    
    


document=Document()
document.add_picture('profile.jpg',width=Inches(2.0))
name=input('Enter your name')
say('Hello' + name + 'how are u?')
say('Whats your phone number?')
contact=input('Enter contact no')
say('Whats your mail id?')
email=input('Enter email id')

document.add_paragraph(name + " " + "|" + " "+ contact + " " + "|" + " " + email )


#About Me Section
document.add_heading('About Me')
document.add_paragraph(input('Enter bio'))

#Work Experience Section
document.add_heading('Work Experience')
p=document.add_paragraph()
company=input('Enter company name')
start_date=input('Enter start date')
end_date=input('Enter end date')
exp=input('Enter your experience')
p.add_run('Name of Company:' + " " + company + "\n").bold=True
p.add_run('Start Date:' + " " + start_date + "-" + "End Date:" + " " + end_date + "\n" ).italic=True
p.add_run("Experience:" + exp).italic=True

while True:
    has_more_exp=input('Do u want to add more experiences? yes/no?')
    if(has_more_exp.lower()=='yes'):
        company=input('Enter company name')
        start_date=input('Enter start date')
        end_date=input('Enter end date')
        exp=input('Enter your experience')
        p.add_run('Name of Company:' + " " + company + "\n").bold=True
        p.add_run('Start Date:' + " " + start_date + "-" + "End Date:" + " " + end_date + "\n" ).italic=True
        p.add_run("Experience:" + exp).italic=True
    else:
        break
    
#Education Section
document.add_heading('Education')
b=document.add_paragraph()
sc_name=input('Enter school name')
from_date=input('Enter sart date')
end_date=input('Enter end date')
percentage=input('Enter percentege/CGPA')
b.add_run(sc_name + "\n").bold=True
b.add_run(from_date + "-" + end_date + "\n").italic=True
b.add_run(percentage).bold=True
while True:
    has_more_ed=input('Want to add more schools? yes/no')
    if(has_more_ed.lower()=='yes'):
        b=document.add_paragraph()
        sc_name=input('Enter school name')
        from_date=input('Enter sart date')
        end_date=input('Enter end date')
        percentage=input('Enter percentege/CGPA')
        b.add_run(sc_name + "\n").bold=True
        b.add_run(from_date + "-" + end_date + "\n").italic=True
        b.add_run(percentage).bold=True
    else:
        break
                        


#Skills Section
document.add_heading('Skills')
a=document.add_paragraph()            
skill_name=input('Enter skill name')
proficiency=input('Enter proficiency')
a.add_run(skill_name + "-" + proficiency).bold=True
while True:
    has_more_skills=input('Do you want to add more skills? yes/no')
    if(has_more_skills.lower()=='yes'):
        a=document.add_paragraph()            
        skill_name=input('Enter skill name')
        proficiency=input('Enter proficiency')
        a.add_run(skill_name + "-" + proficiency).bold=True
        a.style='List Bullet'
    else:
        break
    
    
#Achievements Section
document.add_heading('Achievements')
c=document.add_paragraph()
ac_name=input('Enter achievement name')
desc=input('Add description')
c.add_run(ac_name + "\n").bold=True
c.add_run(desc)
while True:
    has_more_ac=input('Do u want to add more? yes/no')
    if(has_more_ac.lower()=='yes'):
        c=document.add_paragraph()
        ac_name=input('Enter achievement name')
        desc=input('Add description')
        c.add_run(ac_name + "\n").bold=True
        c.add_run(desc)
    else:
        break
    
#Hobbies Section
document.add_heading('Hobbies')
h=document.add_paragraph()
hobby=input('Enter hobby')
h.add_run(hobby)
while True:
    has_more_hobbies=input('Do u want to add more? yes/no')
    if(has_more_hobbies.lower()=='yes'):
        h=document.add_paragraph()
        hobby=input('Enter hobby')
        h.add_run(hobby)
    else:
        break
say('Thanks!!! Your resume will is ready')

    

        
        

    


        

document.save('CV.docx')
